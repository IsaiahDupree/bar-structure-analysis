"""
Create Professional Word Document for Problem Set 3 Solutions
"""
import os
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import matplotlib.pyplot as plt
import numpy as np

def create_word_report():
    # Create a new Document
    document = Document()
    
    # Set the document properties (these don't appear in the document itself)
    document.core_properties.title = "Problem Set 3 Solutions"
    
    # Add introduction
    document.add_paragraph('This report presents the solutions to Problem Set 3, focusing on finite element analysis for one-dimensional elements. The problems involve shape functions, displacements, strains, and stiffness matrices for 2-node and 3-node elements.')
    
    # Make sure figures directory exists
    if not os.path.exists('figures'):
        os.makedirs('figures')
    
    # Add Problem 1
    document.add_heading('Problem 1: Three 2-Node Elements', level=1)
    
    # Add problem statement
    p = document.add_paragraph('Consider a structure consisting of three 2-node elements with nodes at x = 0, 2, 4, 6. The displacements at nodes 1, 2, 3, and 4 are 1, 3, 4, and 6 respectively.')
    
    # Add a simple illustration
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    
    # Create a simple figure showing the structure
    fig, ax = plt.subplots(figsize=(8, 2))
    node_positions = [0, 2, 4, 6]
    ax.plot(node_positions, [0, 0, 0, 0], 'ko-', markersize=10)
    
    # Add node labels
    for i, x in enumerate(node_positions):
        ax.text(x, -0.1, f"Node {i+1}", ha='center')
    
    # Add element labels
    for i in range(len(node_positions)-1):
        mid_x = (node_positions[i] + node_positions[i+1]) / 2
        ax.text(mid_x, 0.1, f"Element {i+1}", ha='center')
    
    ax.set_ylim(-0.2, 0.2)
    ax.axis('off')
    plt.tight_layout()
    
    # Save figure for inclusion in the document
    fig_path = os.path.join('figures', 'problem1_structure.png')
    plt.savefig(fig_path, dpi=300)
    plt.close()
    
    # Add the figure to the document
    run.add_picture(fig_path, width=Inches(6))
    
    # Part 1: Shape functions
    document.add_heading('1. Shape Functions', level=2)
    p = document.add_paragraph('For 2-node elements, the shape functions are:')
    p = document.add_paragraph('N₁(ξ) = (1-ξ)/2')
    p = document.add_paragraph('N₂(ξ) = (1+ξ)/2')
    p = document.add_paragraph('where ξ is the local coordinate ranging from -1 to 1 within each element.')
    
    # Part 2: Displacements
    document.add_heading('2. Displacements at ξ = 0.5', level=2)
    p = document.add_paragraph('The displacement at ξ = 0.5 in each element is calculated using the shape functions:')
    p = document.add_paragraph('u(ξ) = N₁(ξ)u₁ + N₂(ξ)u₂')
    p = document.add_paragraph('For ξ = 0.5:')
    p = document.add_paragraph('N₁(0.5) = 0.25')
    p = document.add_paragraph('N₂(0.5) = 0.75')
    
    # Add element calculations
    p = document.add_paragraph('Element 1: u(ξ=0.5) = 0.25 × 1 + 0.75 × 3 = 2.5')
    p = document.add_paragraph('Element 2: u(ξ=0.5) = 0.25 × 3 + 0.75 × 4 = 3.75')
    p = document.add_paragraph('Element 3: u(ξ=0.5) = 0.25 × 4 + 0.75 × 6 = 5.5')
    
    # Add displacement plot
    if os.path.exists(os.path.join('figures', 'ps3_problem1_displacement.png')):
        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(os.path.join('figures', 'ps3_problem1_displacement.png'), width=Inches(6))
        p = document.add_paragraph('Figure 1: Displacement field for the three 2-node elements.')
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Part 3: Strains
    document.add_heading('3. Strains at ξ = 0.5', level=2)
    p = document.add_paragraph('For 2-node elements, the strain is constant throughout the element:')
    p = document.add_paragraph('ε = (u₂ - u₁)/L')
    
    # Add strain calculations
    p = document.add_paragraph('Element 1: ε = (3 - 1)/2 = 1.0')
    p = document.add_paragraph('Element 2: ε = (4 - 3)/2 = 0.5')
    p = document.add_paragraph('Element 3: ε = (6 - 4)/2 = 1.0')
    
    # Add strain plot
    if os.path.exists(os.path.join('figures', 'ps3_problem1_strains.png')):
        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(os.path.join('figures', 'ps3_problem1_strains.png'), width=Inches(6))
        p = document.add_paragraph('Figure 2: Element strains for the three 2-node elements.')
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Part 4: Stiffness Matrices
    document.add_heading('4. Stiffness Matrices', level=2)
    p = document.add_paragraph('For a 2-node element with Young\'s modulus E and cross-sectional area A, the stiffness matrix is:')
    p = document.add_paragraph('k = (E·A/L) × [1, -1; -1, 1]')
    
    # For simplicity, assuming E=A=1
    p = document.add_paragraph('Assuming E = A = 1 for simplicity:')
    p = document.add_paragraph('Element 1 (L=2): k₁ = 0.5 × [1, -1; -1, 1] = [0.5, -0.5; -0.5, 0.5]')
    p = document.add_paragraph('Element 2 (L=2): k₂ = 0.5 × [1, -1; -1, 1] = [0.5, -0.5; -0.5, 0.5]')
    p = document.add_paragraph('Element 3 (L=2): k₃ = 0.5 × [1, -1; -1, 1] = [0.5, -0.5; -0.5, 0.5]')
    
    # Add Problem 2
    document.add_heading('Problem 2: Two 3-Node Elements', level=1)
    
    # Add problem statement
    p = document.add_paragraph('Consider a structure consisting of two 3-node elements with nodes at x = 0, 2, 4, 6, 8. The displacements at nodes 1, 2, 3, 4, and 5 are 0, -1, 2, -1, and 4 respectively. Element 1 connects nodes 1, 2, and 3, while Element 2 connects nodes 3, 4, and 5.')
    
    # Add a simple illustration
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    
    # Create a simple figure showing the structure
    fig, ax = plt.subplots(figsize=(8, 2))
    node_positions = [0, 2, 4, 6, 8]
    ax.plot(node_positions, [0, 0, 0, 0, 0], 'ko-', markersize=10)
    
    # Add node labels
    for i, x in enumerate(node_positions):
        ax.text(x, -0.1, f"Node {i+1}", ha='center')
    
    # Add element labels
    mid_x1 = (node_positions[0] + node_positions[2]) / 2
    mid_x2 = (node_positions[2] + node_positions[4]) / 2
    ax.text(mid_x1, 0.1, "Element 1", ha='center')
    ax.text(mid_x2, 0.1, "Element 2", ha='center')
    
    ax.set_ylim(-0.2, 0.2)
    ax.axis('off')
    plt.tight_layout()
    
    # Save figure for inclusion in the document
    fig_path = os.path.join('figures', 'problem2_structure.png')
    plt.savefig(fig_path, dpi=300)
    plt.close()
    
    # Add the figure to the document
    run.add_picture(fig_path, width=Inches(6))
    
    # Part 1: Shape functions
    document.add_heading('1. Shape Functions', level=2)
    p = document.add_paragraph('For 3-node elements, the shape functions are:')
    p = document.add_paragraph('N₁(ξ) = ξ(ξ-1)/2')
    p = document.add_paragraph('N₂(ξ) = (1+ξ)(1-ξ)')
    p = document.add_paragraph('N₃(ξ) = ξ(ξ+1)/2')
    p = document.add_paragraph('where ξ is the local coordinate ranging from -1 to 1 within each element, with ξ = -1, 0, 1 corresponding to the three nodes.')
    
    # Part 2: Displacements
    document.add_heading('2. Displacements at ξ = -0.5', level=2)
    p = document.add_paragraph('The displacement at ξ = -0.5 in each element is calculated using the shape functions:')
    p = document.add_paragraph('u(ξ) = N₁(ξ)u₁ + N₂(ξ)u₂ + N₃(ξ)u₃')
    p = document.add_paragraph('For ξ = -0.5:')
    p = document.add_paragraph('N₁(-0.5) = 0.125')
    p = document.add_paragraph('N₂(-0.5) = 0.75')
    p = document.add_paragraph('N₃(-0.5) = -0.125')
    
    # Add element calculations
    p = document.add_paragraph('Element 1: u(ξ=-0.5) = 0.125 × 0 + 0.75 × (-1) + (-0.125) × 2 = -1.0')
    p = document.add_paragraph('Element 2: u(ξ=-0.5) = 0.125 × 2 + 0.75 × (-1) + (-0.125) × 4 = -0.75')
    
    # Add displacement plot
    if os.path.exists(os.path.join('figures', 'ps3_problem2_displacement.png')):
        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(os.path.join('figures', 'ps3_problem2_displacement.png'), width=Inches(6))
        p = document.add_paragraph('Figure 3: Displacement field for the two 3-node elements.')
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Part 3: Strains
    document.add_heading('3. Strains at ξ = -0.5', level=2)
    p = document.add_paragraph('For 3-node elements, the strain varies within the element and is calculated using:')
    p = document.add_paragraph('ε = B·u, where B = [dN₁/dx, dN₂/dx, dN₃/dx]')
    p = document.add_paragraph('At ξ = -0.5:')
    
    # Derivatives of shape functions
    p = document.add_paragraph('dN₁/dξ = ξ - 0.5 = -1.0')
    p = document.add_paragraph('dN₂/dξ = -2ξ = 1.0')
    p = document.add_paragraph('dN₃/dξ = ξ + 0.5 = 0.0')
    
    # Jacobian
    p = document.add_paragraph('For Element 1 (L=4): J = L/2 = 2')
    p = document.add_paragraph('For Element 2 (L=4): J = L/2 = 2')
    
    # Shape function derivatives with respect to x
    p = document.add_paragraph('dN₁/dx = dN₁/dξ · dξ/dx = (-1.0)/2 = -0.5')
    p = document.add_paragraph('dN₂/dx = dN₂/dξ · dξ/dx = (1.0)/2 = 0.5')
    p = document.add_paragraph('dN₃/dx = dN₃/dξ · dξ/dx = (0.0)/2 = 0.0')
    
    # Strain calculations
    p = document.add_paragraph('Element 1: ε = -0.5 × 0 + 0.5 × (-1) + 0.0 × 2 = -0.5')
    p = document.add_paragraph('Element 2: ε = -0.5 × 2 + 0.5 × (-1) + 0.0 × 4 = -1.5')
    
    # Add strain plot
    if os.path.exists(os.path.join('figures', 'ps3_problem2_strains.png')):
        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(os.path.join('figures', 'ps3_problem2_strains.png'), width=Inches(6))
        p = document.add_paragraph('Figure 4: Element strains for the two 3-node elements.')
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Part 4: Stiffness Matrices
    document.add_heading('4. Stiffness Matrices', level=2)
    p = document.add_paragraph('For a 3-node element, the stiffness matrix is calculated using numerical integration:')
    p = document.add_paragraph('k = ∫(B^T·E·B·A·det(J))dξ')
    p = document.add_paragraph('Using integration points ξ = -0.5 and ξ = 0.5 with equal weights of 1.0:')
    
    # Stiffness matrices visualization
    if os.path.exists(os.path.join('figures', 'ps3_problem2_stiffness.png')):
        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(os.path.join('figures', 'ps3_problem2_stiffness.png'), width=Inches(6))
        p = document.add_paragraph('Figure 5: Stiffness matrices for the two 3-node elements.')
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Conclusion
    document.add_heading('Conclusion', level=1)
    p = document.add_paragraph('This report has presented solutions for two finite element problems involving 2-node and 3-node elements. We derived the shape functions, calculated displacements and strains at specific points in the elements, and computed element stiffness matrices using appropriate integration methods. The results demonstrate how the finite element method can effectively model displacement fields and strain distributions in one-dimensional structures.')
    
    # Save the document
    document.save('Problem_Set_3_Solutions.docx')
    print("Word document created: Problem_Set_3_Solutions.docx")

if __name__ == "__main__":
    create_word_report()
