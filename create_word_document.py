"""
Script to create a Word document containing the code and report
"""
import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_word_document():
    # Create a new Document
    doc = Document()
    
    # Add title
    title = doc.add_heading('Composite Bar Structure Analysis using Finite Element Method', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add introduction
    doc.add_heading('Introduction', level=1)
    intro = doc.add_paragraph()
    intro.add_run('This document presents a finite element method (FEM) implementation for analyzing a composite bar structure. ')
    intro.add_run('The bar consists of three segments with different cross-sectional areas and material properties. ')
    intro.add_run('Both analytical and numerical solutions are presented, along with error analysis and visualization.')
    
    # Add problem description
    doc.add_heading('Problem Description', level=1)
    problem = doc.add_paragraph()
    problem.add_run('A bar structure with the following properties is analyzed:').bold = True
    
    # Add bullet points for problem description
    bullets = [
        'The bar consists of three segments, each of length L = 500 mm.',
        'Segment 1: Cross-sectional area varies linearly from A₁ (200 mm²) to A₂ (100 mm²), Modulus E₁ = 130 GPa',
        'Segment 2: Constant cross-sectional area A₂ (100 mm²), Modulus E₁ = 130 GPa',
        'Segment 3: Constant cross-sectional area A₃ (50 mm²), Modulus E₂ = 200 GPa',
        'Forces applied: F₁ = 20 kN, F₂ = 40 kN, F₃ = 20 kN at the ends of segments 1, 2, and 3 respectively',
        'The left end of the bar is fixed (displacement = 0)'
    ]
    
    for bullet in bullets:
        bullet_para = doc.add_paragraph(bullet, style='List Bullet')
    
    # Add methodology
    doc.add_heading('Methodology', level=1)
    
    doc.add_heading('Analytical Solution', level=2)
    analytical = doc.add_paragraph()
    analytical.add_run('For a bar with varying cross-section, the stress is given by σ(x) = F/A(x), ')
    analytical.add_run('where F is the internal force and A(x) is the cross-sectional area at position x. ')
    analytical.add_run('The displacement is calculated by integrating u(x) = ∫[F/(A(x)·E(x))]dx, ')
    analytical.add_run('where E(x) is the elastic modulus at position x.')
    
    doc.add_heading('Finite Element Method (FEM)', level=2)
    fem = doc.add_paragraph()
    fem.add_run('The finite element method implementation follows these steps:').bold = True
    
    fem_steps = [
        'Partition the structure into a number of elements',
        'Use 2-node linear elements with one degree of freedom per node',
        'Calculate element stiffness matrices: k = (A·E/L) * [1 -1; -1 1]',
        'Assemble the global stiffness matrix and force vector',
        'Apply boundary conditions (fixed left end)',
        'Solve the system KU = F for nodal displacements',
        'Calculate element stresses and compare with analytical solution',
        'Refine the mesh until the error is below 5%'
    ]
    
    for step in fem_steps:
        step_para = doc.add_paragraph(step, style='List Number')
    
    # Add Python code section
    doc.add_heading('Python Implementation', level=1)
    
    # Add main.py code
    doc.add_heading('main.py', level=2)
    with open('main.py', 'r') as f:
        code = f.read()
    
    main_code = doc.add_paragraph()
    main_code.add_run(code).font.name = 'Courier New'
    main_code.add_run().font.size = Pt(9)
    
    # Add bar_analysis.py code
    doc.add_heading('bar_analysis.py', level=2)
    with open('bar_analysis.py', 'r') as f:
        code = f.read()
    
    analysis_code = doc.add_paragraph()
    analysis_code.add_run(code).font.name = 'Courier New'
    analysis_code.add_run().font.size = Pt(9)
    
    # Add utils.py code
    doc.add_heading('utils.py', level=2)
    with open('utils.py', 'r') as f:
        code = f.read()
    
    utils_code = doc.add_paragraph()
    utils_code.add_run(code).font.name = 'Courier New'
    utils_code.add_run().font.size = Pt(9)
    
    # Add results section
    doc.add_heading('Results and Analysis', level=1)
    
    results = doc.add_paragraph()
    results.add_run('The finite element analysis produces the following results:')
    
    # List expected outputs
    result_items = [
        'Displacement field showing how the bar deforms under the applied loads',
        'Stress distribution along the bar',
        'Error analysis comparing FEM results with analytical solution',
        'Convergence study showing how the error decreases with mesh refinement'
    ]
    
    for item in result_items:
        item_para = doc.add_paragraph(item, style='List Bullet')
    
    # Add analysis insights
    doc.add_heading('Analysis Insights', level=2)
    
    insights = doc.add_paragraph()
    insights.add_run('Key observations from the analysis:').bold = True
    
    insight_items = [
        'Stress discontinuities occur at segment boundaries due to changes in cross-sectional area',
        'Higher stresses are observed in segments with smaller cross-sectional areas',
        'The maximum displacement occurs at the right end of the bar',
        'FEM errors are typically highest at segment transitions',
        'Mesh refinement significantly improves accuracy, particularly at stress concentrations'
    ]
    
    for item in insight_items:
        item_para = doc.add_paragraph(item, style='List Bullet')
    
    # Add conclusion
    doc.add_heading('Conclusion', level=1)
    
    conclusion = doc.add_paragraph()
    conclusion.add_run('This implementation demonstrates the application of the finite element method to analyze a composite bar structure. ')
    conclusion.add_run('The code successfully calculates stresses and displacements, and the mesh refinement strategy ensures that errors remain below the specified threshold of 5%. ')
    conclusion.add_run('The Python implementation provides a flexible framework that can be extended to more complex problems.')
    
    # Save the document
    doc_path = 'Composite_Bar_Structure_Analysis.docx'
    doc.save(doc_path)
    
    print(f"Word document created: {doc_path}")
    return doc_path

if __name__ == "__main__":
    # Ensure we're in the correct directory
    script_dir = os.path.dirname(os.path.abspath(__file__))
    os.chdir(script_dir)
    
    try:
        create_word_document()
    except Exception as e:
        print(f"Error creating Word document: {e}")
        print("You might need to install the python-docx package:")
        print("pip install python-docx")
