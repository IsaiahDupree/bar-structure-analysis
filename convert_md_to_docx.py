"""
Convert Markdown solution file to Word document
"""

import os
from datetime import datetime
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def convert_markdown_to_word(md_path, output_path):
    """
    Convert a Markdown file to a Word document with proper formatting
    """
    # Create a new Word document
    doc = Document()
    
    # Set document properties
    doc.core_properties.title = "Composite Bar Structure Analysis Solution"
    doc.core_properties.author = "Engineering Analysis Team"
    doc.core_properties.created = datetime.now()
    
    # Configure document formatting
    sections = doc.sections
    for section in sections:
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
    
    # Read the Markdown file
    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()
    
    # Process lines and convert to Word format
    in_list = False
    for line in lines:
        line = line.rstrip()
        
        # Blank line
        if not line:
            doc.add_paragraph()
            in_list = False
            continue
        
        # Headings
        if line.startswith('# '):
            p = doc.add_heading(line[2:], level=0)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            continue
        
        if line.startswith('## '):
            p = doc.add_heading(line[3:], level=1)
            continue
            
        if line.startswith('### '):
            p = doc.add_heading(line[4:], level=2)
            continue
        
        # Lists
        if line.startswith('- '):
            if not in_list:
                p = doc.add_paragraph()
                in_list = True
            p = doc.add_paragraph(line[2:], style='List Bullet')
            continue
            
        if line.startswith('  - '):
            p = doc.add_paragraph(line[4:], style='List Bullet 2')
            continue
            
        if line[0].isdigit() and line[1] == '.' and line[2] == ' ':
            if not in_list:
                p = doc.add_paragraph()
                in_list = True
            p = doc.add_paragraph(line[3:], style='List Number')
            continue
        
        # Regular paragraph
        in_list = False
        p = doc.add_paragraph(line)
    
    # Save the document
    doc.save(output_path)
    print(f"Word document created: {output_path}")

if __name__ == "__main__":
    base_dir = os.path.dirname(os.path.abspath(__file__))
    
    # Input markdown file
    md_file = os.path.join(base_dir, "Composite_Bar_Structure_Analysis_Solution.md")
    
    # Output docx file with timestamp
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    docx_file = os.path.join(base_dir, f"Composite_Bar_Structure_Analysis_Solution_{timestamp}.docx")
    
    # Convert to Word
    convert_markdown_to_word(md_file, docx_file)
