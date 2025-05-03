"""
Create Simplified Word Document for Problem Set 3 Solutions
This script generates a simplified report that uses the original problem images
and emphasizes the mathematical work with fewer graphs.
"""
import os
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import numpy as np

def create_simplified_word_report():
    # Create a new Document
    document = Document()
    
    # Set the document properties
    document.core_properties.title = "Problem Set 3 Solutions - Simplified"
    document.core_properties.author = "Student"
    
    # Add a title
    title = document.add_heading('Problem Set 3: Finite Element Analysis (Simplified Version)', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add problem statements section
    document.add_heading('Problem Statements', level=1)
    p = document.add_paragraph('Below are the original problem statements for Problem Set 3:')
    
    # Add the original problem images
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    try:
        run.add_picture("IMG_5525.jpeg", width=Inches(6))
        p = document.add_paragraph("Figure 1: Problem 1 Statement")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    except:
        p = document.add_paragraph("Note: Original problem image 'IMG_5525.jpeg' not found. Please insert it manually.")
    
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    try:
        run.add_picture("IMG_5526.jpeg", width=Inches(6))
        p = document.add_paragraph("Figure 2: Problem 2 Statement")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    except:
        p = document.add_paragraph("Note: Original problem image 'IMG_5526.jpeg' not found. Please insert it manually.")
    
    # Problem 1
    document.add_heading('Problem 1: Three 2-Node Elements', level=1)
    
    # Part 1: Shape Functions
    document.add_heading('1. Shape Functions', level=2)
    p = document.add_paragraph('For 2-node elements, the shape functions are:')
    p = document.add_paragraph('N₁(ξ) = (1-ξ)/2')
    p = document.add_paragraph('N₂(ξ) = (1+ξ)/2')
    p = document.add_paragraph('where ξ is the local coordinate ranging from -1 to 1 within each element.')
    
    # Part 2: Displacements
    document.add_heading('2. Displacements at ξ = 0.5', level=2)
    p = document.add_paragraph('The displacement at ξ = 0.5 in each element is calculated using the shape functions:')
    p = document.add_paragraph('u(ξ) = N₁(ξ)u₁ + N₂(ξ)u₂')
    p = document.add_paragraph('For ξ = 0.5:')
    p = document.add_paragraph('N₁(0.5) = (1-0.5)/2 = 0.25')
    p = document.add_paragraph('N₂(0.5) = (1+0.5)/2 = 0.75')
    
    # Add element calculations
    p = document.add_paragraph('Element 1: u(ξ=0.5) = 0.25 × 1 + 0.75 × 3 = 0.25 + 2.25 = 2.5')
    p = document.add_paragraph('Element 2: u(ξ=0.5) = 0.25 × 3 + 0.75 × 4 = 0.75 + 3.0 = 3.75')
    p = document.add_paragraph('Element 3: u(ξ=0.5) = 0.25 × 4 + 0.75 × 6 = 1.0 + 4.5 = 5.5')
    
    # Part 3: Strains
    document.add_heading('3. Strains at ξ = 0.5', level=2)
    p = document.add_paragraph('For 2-node elements, the strain is constant throughout the element:')
    p = document.add_paragraph('ε = (u₂ - u₁)/L')
    
    # Add strain calculations
    p = document.add_paragraph('Element 1: ε = (3 - 1)/2 = 1.0')
    p = document.add_paragraph('Element 2: ε = (4 - 3)/2 = 0.5')
    p = document.add_paragraph('Element 3: ε = (6 - 4)/2 = 1.0')
    
    # Part 4: Stiffness Matrices
    document.add_heading('4. Stiffness Matrices', level=2)
    p = document.add_paragraph('For a 2-node element with Young\'s modulus E and cross-sectional area A, the stiffness matrix is:')
    p = document.add_paragraph('k = (E·A/L) × [1, -1; -1, 1]')
    
    # For simplicity, assuming E=A=1
    p = document.add_paragraph('Assuming E = A = 1 for simplicity:')
    p = document.add_paragraph('Element 1 (L=2): k₁ = 0.5 × [1, -1; -1, 1] = [0.5, -0.5; -0.5, 0.5]')
    p = document.add_paragraph('Element 2 (L=2): k₂ = 0.5 × [1, -1; -1, 1] = [0.5, -0.5; -0.5, 0.5]')
    p = document.add_paragraph('Element 3 (L=2): k₃ = 0.5 × [1, -1; -1, 1] = [0.5, -0.5; -0.5, 0.5]')
    
    # Problem 2
    document.add_heading('Problem 2: Two 3-Node Elements', level=1)
    
    # Part 1: Shape Functions
    document.add_heading('1. Shape Functions', level=2)
    p = document.add_paragraph('For 3-node elements, the shape functions are:')
    p = document.add_paragraph('N₁(ξ) = ξ(ξ-1)/2')
    p = document.add_paragraph('N₂(ξ) = (1+ξ)(1-ξ)')
    p = document.add_paragraph('N₃(ξ) = ξ(ξ+1)/2')
    p = document.add_paragraph('where ξ is the local coordinate ranging from -1 to 1 within each element, with ξ = -1, 0, 1 corresponding to the three nodes.')
    
    # Part 2: Displacements
    document.add_heading('2. Displacements at ξ = -0.5', level=2)
    p = document.add_paragraph('The displacement at ξ = -0.5 in each element is calculated using the shape functions:')
    p = document.add_paragraph('u(ξ) = N₁(ξ)u₁ + N₂(ξ)u₂ + N₃(ξ)u₃')
    p = document.add_paragraph('For ξ = -0.5:')
    p = document.add_paragraph('N₁(-0.5) = (-0.5)(-0.5-1)/2 = (-0.5)(-1.5)/2 = 0.75/2 = 0.125')
    p = document.add_paragraph('N₂(-0.5) = (1+(-0.5))(1-(-0.5)) = 0.5 × 1.5 = 0.75')
    p = document.add_paragraph('N₃(-0.5) = (-0.5)(-0.5+1)/2 = (-0.5)(0.5)/2 = -0.125')
    
    # Add element calculations
    p = document.add_paragraph('Element 1: u(ξ=-0.5) = 0.125 × 0 + 0.75 × (-1) + (-0.125) × 2 = 0 - 0.75 - 0.25 = -1.0')
    p = document.add_paragraph('Element 2: u(ξ=-0.5) = 0.125 × 2 + 0.75 × (-1) + (-0.125) × 4 = 0.25 - 0.75 - 0.5 = -1.0')
    
    # Part 3: Strains
    document.add_heading('3. Strains at ξ = -0.5', level=2)
    p = document.add_paragraph('For 3-node elements, the strain varies within the element and is calculated using:')
    p = document.add_paragraph('ε = B · u, where B = [dN₁/dx, dN₂/dx, dN₃/dx]')
    p = document.add_paragraph('At ξ = -0.5:')
    
    # Derivatives of shape functions
    p = document.add_paragraph('dN₁/dξ = ξ - 0.5 = -0.5 - 0.5 = -1.0')
    p = document.add_paragraph('dN₂/dξ = -2ξ = -2 × (-0.5) = 1.0')
    p = document.add_paragraph('dN₃/dξ = ξ + 0.5 = -0.5 + 0.5 = 0.0')
    
    # Jacobian
    p = document.add_paragraph('For Element 1 (L=4): J = L/2 = 2')
    p = document.add_paragraph('For Element 2 (L=4): J = L/2 = 2')
    
    # Shape function derivatives with respect to x
    p = document.add_paragraph('dN₁/dx = dN₁/dξ · dξ/dx = (-1.0)/2 = -0.5')
    p = document.add_paragraph('dN₂/dx = dN₂/dξ · dξ/dx = (1.0)/2 = 0.5')
    p = document.add_paragraph('dN₃/dx = dN₃/dξ · dξ/dx = (0.0)/2 = 0.0')
    
    # Strain calculations
    p = document.add_paragraph('Element 1: ε = -0.5 × 0 + 0.5 × (-1) + 0.0 × 2 = 0 - 0.5 + 0 = -0.5')
    p = document.add_paragraph('Element 2: ε = -0.5 × 2 + 0.5 × (-1) + 0.0 × 4 = -1.0 - 0.5 + 0 = -1.5')
    
    # Part 4: Stiffness Matrices
    document.add_heading('4. Stiffness Matrices', level=2)
    p = document.add_paragraph('For a 3-node element, the stiffness matrix is calculated using numerical integration:')
    p = document.add_paragraph('k = ∫(B^T·E·B·A·det(J))dξ')
    p = document.add_paragraph('Using integration points ξ = -0.5 and ξ = 0.5 with equal weights of 1.0 and assuming E = A = 1:')
    
    # Add stiffness matrix calculations
    p = document.add_paragraph('For Element 1:')
    p = document.add_paragraph('k₁ = [0.5, -0.5, 0; -0.5, 1.0, -0.5; 0, -0.5, 0.5]')
    
    p = document.add_paragraph('For Element 2:')
    p = document.add_paragraph('k₂ = [0.5, -0.5, 0; -0.5, 1.0, -0.5; 0, -0.5, 0.5]')
    
    # Summary section
    document.add_heading('Summary of Results', level=1)
    
    # Problem 1 summary
    document.add_heading('Problem 1 Results', level=2)
    
    # Create a table for Problem 1 results
    table = document.add_table(rows=3, cols=4)
    table.style = 'Table Grid'
    
    # Add headers
    header_cells = table.rows[0].cells
    header_cells[0].text = 'Quantity'
    header_cells[1].text = 'Element 1'
    header_cells[2].text = 'Element 2'
    header_cells[3].text = 'Element 3'
    
    # Add displacement row
    disp_cells = table.rows[1].cells
    disp_cells[0].text = 'Displacement at ξ = 0.5'
    disp_cells[1].text = '2.5'
    disp_cells[2].text = '3.75'
    disp_cells[3].text = '5.5'
    
    # Add strain row
    strain_cells = table.rows[2].cells
    strain_cells[0].text = 'Strain at ξ = 0.5'
    strain_cells[1].text = '1.0'
    strain_cells[2].text = '0.5'
    strain_cells[3].text = '1.0'
    
    # Problem 2 summary
    document.add_heading('Problem 2 Results', level=2)
    
    # Create a table for Problem 2 results
    table = document.add_table(rows=3, cols=3)
    table.style = 'Table Grid'
    
    # Add headers
    header_cells = table.rows[0].cells
    header_cells[0].text = 'Quantity'
    header_cells[1].text = 'Element 1'
    header_cells[2].text = 'Element 2'
    
    # Add displacement row
    disp_cells = table.rows[1].cells
    disp_cells[0].text = 'Displacement at ξ = -0.5'
    disp_cells[1].text = '-1.0'
    disp_cells[2].text = '-1.0'
    
    # Add strain row
    strain_cells = table.rows[2].cells
    strain_cells[0].text = 'Strain at ξ = -0.5'
    strain_cells[1].text = '-0.5'
    strain_cells[2].text = '-1.5'
    
    # Save the document
    document.save('Problem_Set_3_Simplified_Solutions.docx')
    print("Simplified Word document created: Problem_Set_3_Simplified_Solutions.docx")

if __name__ == "__main__":
    create_simplified_word_report()
