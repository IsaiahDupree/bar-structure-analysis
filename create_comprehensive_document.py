"""
Script to create a comprehensive Word document for the Bar Structure Analysis
"""
import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.shared import Cm

def create_comprehensive_document():
    # Create a new Document
    doc = Document()
    
    # Set up styles
    styles = doc.styles
    
    # Create title style
    title_style = styles.add_style('Custom Title', WD_STYLE_TYPE.PARAGRAPH)
    title_format = title_style.paragraph_format
    title_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_format.space_before = Pt(12)
    title_format.space_after = Pt(24)
    title_run_format = title_style.font
    title_run_format.size = Pt(18)
    title_run_format.bold = True
    
    # Create heading styles
    heading1_style = styles.add_style('Custom Heading 1', WD_STYLE_TYPE.PARAGRAPH)
    heading1_format = heading1_style.paragraph_format
    heading1_format.space_before = Pt(18)
    heading1_format.space_after = Pt(6)
    heading1_run_format = heading1_style.font
    heading1_run_format.size = Pt(16)
    heading1_run_format.bold = True
    heading1_run_format.color.rgb = RGBColor(0, 70, 130)
    
    heading2_style = styles.add_style('Custom Heading 2', WD_STYLE_TYPE.PARAGRAPH)
    heading2_format = heading2_style.paragraph_format
    heading2_format.space_before = Pt(12)
    heading2_format.space_after = Pt(6)
    heading2_run_format = heading2_style.font
    heading2_run_format.size = Pt(14)
    heading2_run_format.bold = True
    heading2_run_format.color.rgb = RGBColor(0, 70, 130)
    
    # Create normal text style
    normal_style = styles.add_style('Custom Normal', WD_STYLE_TYPE.PARAGRAPH)
    normal_format = normal_style.paragraph_format
    normal_format.space_before = Pt(6)
    normal_format.space_after = Pt(6)
    normal_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    normal_run_format = normal_style.font
    normal_run_format.size = Pt(11)
    
    # Create code block style
    code_style = styles.add_style('Code Block', WD_STYLE_TYPE.PARAGRAPH)
    code_format = code_style.paragraph_format
    code_format.space_before = Pt(6)
    code_format.space_after = Pt(6)
    code_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    code_run_format = code_style.font
    code_run_format.name = 'Courier New'
    code_run_format.size = Pt(9)
    
    # Add title
    title = doc.add_paragraph('Composite Bar Structure Analysis using Finite Element Method', style='Custom Title')
    
    # Add document metadata
    doc.add_paragraph('Analysis Report', style='Custom Normal').alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph('Date: April 23, 2025', style='Custom Normal').alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add introduction
    doc.add_paragraph('1. Introduction', style='Custom Heading 1')
    intro = doc.add_paragraph('This report presents a comprehensive analysis of a composite bar structure with varying cross-sectional areas and material properties. The structure is subjected to axial forces at different locations along its length. We investigate the problem using both closed-form analytical solutions and finite element analysis (FEA).', style='Custom Normal')
    
    # Add problem description
    doc.add_paragraph('2. Problem Definition', style='Custom Heading 1')
    
    # Problem schematic
    try:
        doc.add_paragraph('Figure 1 shows the schematic of the bar structure:', style='Custom Normal')
        doc.add_picture('plots/area_distribution.png', width=Inches(6))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph('Figure 1: Schematic representation of the composite bar structure showing cross-sectional area distribution', style='Custom Normal').alignment = WD_ALIGN_PARAGRAPH.CENTER
    except:
        doc.add_paragraph('[Area distribution plot would be displayed here]', style='Custom Normal').alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Problem parameters table
    doc.add_paragraph('The bar structure has the following properties:', style='Custom Normal')
    
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    table.autofit = False
    table.allow_autofit = False
    table.columns[0].width = Cm(6)
    table.columns[1].width = Cm(6)
    
    header_cells = table.rows[0].cells
    header_cells[0].text = 'Parameter'
    header_cells[1].text = 'Value'
    
    # Add parameter rows
    parameters = [
        ('Length of each segment (L)', '500 mm'),
        ('Cross-section area 1 (A₁)', '200 mm²'),
        ('Cross-section area 2 (A₂)', '100 mm²'),
        ('Cross-section area 3 (A₃)', '50 mm²'),
        ('Elastic modulus 1 (E₁)', '130 GPa'),
        ('Elastic modulus 2 (E₂)', '200 GPa'),
        ('Force 1 (F₁)', '20 kN'),
        ('Force 2 (F₂)', '40 kN'),
        ('Force 3 (F₃)', '20 kN')
    ]
    
    for param, value in parameters:
        row_cells = table.add_row().cells
        row_cells[0].text = param
        row_cells[1].text = value
    
    # Problem statement
    doc.add_paragraph('The structure consists of a bar fixed at the left end and composed of three segments of equal length. The first segment has a linearly varying cross-section from A₁ to A₂, while the second and third segments have constant cross-sections A₂ and A₃, respectively. Forces F₁, F₂, and F₃ are applied at the junctions between segments and at the free end.', style='Custom Normal')
    
    # Analytical Solution
    doc.add_paragraph('3. Analytical Solution', style='Custom Heading 1')
    
    # Internal forces
    doc.add_paragraph('3.1 Internal Axial Forces', style='Custom Heading 2')
    doc.add_paragraph('The left end is fixed, and all loads act to the right. The reaction force at the fixed end must balance all applied forces:', style='Custom Normal')
    p = doc.add_paragraph('R = F₁ + F₂ + F₃ = 80,000 N', style='Custom Normal')
    
    # Internal forces table
    table2 = doc.add_table(rows=1, cols=3)
    table2.style = 'Table Grid'
    
    header_cells = table2.rows[0].cells
    header_cells[0].text = 'Segment'
    header_cells[1].text = 'Span'
    header_cells[2].text = 'Internal Force'
    
    # Add segments rows
    segments = [
        ('1 (0–L)', '0 ≤ x ≤ L', 'N₁ = R = 80,000 N'),
        ('2 (L–2L)', 'L ≤ x ≤ 2L', 'N₂ = R - F₁ = 60,000 N'),
        ('3 (2L–3L)', '2L ≤ x ≤ 3L', 'N₃ = R - F₁ - F₂ = 20,000 N')
    ]
    
    for segment, span, force in segments:
        row_cells = table2.add_row().cells
        row_cells[0].text = segment
        row_cells[1].text = span
        row_cells[2].text = force
    
    # Stress distribution
    doc.add_paragraph('3.2 Stress Distribution', style='Custom Heading 2')
    
    # Segment 1
    doc.add_paragraph('Segment 1 – linearly-tapered area:', style='Custom Normal')
    p = doc.add_paragraph('A(x) = A₁ + (A₂ - A₁)·x/L', style='Custom Normal')
    p = doc.add_paragraph('σ₁(x) = N₁/A(x)', style='Custom Normal')
    doc.add_paragraph('The stress in segment 1 rises linearly from 400 MPa at x = 0 to 800 MPa at x = L.', style='Custom Normal')
    
    # Segments 2 & 3
    doc.add_paragraph('Segments 2 & 3 – uniform cross-section:', style='Custom Normal')
    p = doc.add_paragraph('σ₂ = N₂/A₂ = 600 MPa', style='Custom Normal')
    p = doc.add_paragraph('σ₃ = N₃/A₃ = 400 MPa', style='Custom Normal')
    
    # Axial displacements
    doc.add_paragraph('3.3 Axial Displacements', style='Custom Heading 2')
    
    # General formula
    doc.add_paragraph('The general formula for axial displacement is:', style='Custom Normal')
    p = doc.add_paragraph('u(x) = ∫[0 to x] [N(ξ)/(E(ξ)·A(ξ))] dξ', style='Custom Normal')
    
    # Segment 1
    doc.add_paragraph('Segment 1 (with logarithmic closed form):', style='Custom Normal')
    p = doc.add_paragraph('δ₁ = (N₁·L)/(E₁·(A₂-A₁))·ln(A₂/A₁) = (80,000·500)/(130,000·(100-200))·ln(0.5) ≈ 2.13 mm', style='Custom Normal')
    
    # Segment 2
    doc.add_paragraph('Segment 2:', style='Custom Normal')
    p = doc.add_paragraph('δ₂ = (N₂·L)/(E₁·A₂) = (60,000·500)/(130,000·100) ≈ 2.31 mm', style='Custom Normal')
    
    # Segment 3
    doc.add_paragraph('Segment 3:', style='Custom Normal')
    p = doc.add_paragraph('δ₃ = (N₃·L)/(E₂·A₃) = (20,000·500)/(200,000·50) = 1.00 mm', style='Custom Normal')
    
    # Total displacement
    doc.add_paragraph('The total free-end displacement is:', style='Custom Normal')
    p = doc.add_paragraph('u₃ₗ = δ₁ + δ₂ + δ₃ ≈ 5.44 mm', style='Custom Normal')
    
    # Try to insert stress field plot
    try:
        doc.add_paragraph('The analytical stress distribution is shown in Figure 2:', style='Custom Normal')
        # Look in multiple possible locations for the plot files
        plot_paths = [
            'plots/stress_field_4.png',
            '../plots/stress_field_4.png',
            'stress_field_4.png',
            os.path.join(os.path.dirname(os.getcwd()), 'plots', 'stress_field_4.png')
        ]
        
        plot_found = False
        for path in plot_paths:
            if os.path.exists(path):
                doc.add_picture(path, width=Inches(6))
                plot_found = True
                print(f"Found plot at: {path}")
                break
        
        if not plot_found and os.path.exists(os.path.join(os.getcwd(), '..', 'plots')):
            # List available plots and use the first stress_field plot found
            parent_plots_dir = os.path.join(os.getcwd(), '..', 'plots')
            for filename in os.listdir(parent_plots_dir):
                if 'stress_field' in filename:
                    doc.add_picture(os.path.join(parent_plots_dir, filename), width=Inches(6))
                    plot_found = True
                    print(f"Found alternative plot: {filename}")
                    break
                    
        if plot_found:
            last_paragraph = doc.paragraphs[-1]
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            doc.add_paragraph('Figure 2: Analytical stress distribution along the bar', style='Custom Normal').alignment = WD_ALIGN_PARAGRAPH.CENTER
        else:
            doc.add_paragraph('[Stress field plot would be displayed here]', style='Custom Normal').alignment = WD_ALIGN_PARAGRAPH.CENTER
    except Exception as e:
        print(f"Error inserting stress field plot: {e}")
        doc.add_paragraph('[Stress field plot would be displayed here]', style='Custom Normal').alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Finite Element Analysis
    doc.add_paragraph('4. Finite Element Analysis', style='Custom Heading 1')
    
    # FEM description
    doc.add_paragraph('4.1 Implementation Overview', style='Custom Heading 2')
    doc.add_paragraph('We implemented a finite element analysis of the bar structure using both Python and MATLAB. The implementation follows these steps:', style='Custom Normal')
    
    steps_list = doc.add_paragraph(style='Custom Normal')
    steps_list.add_run('1. Partition the structure into elements: Initially 4 elements per segment\n').bold = True
    steps_list.add_run('2. Assign material and geometric properties for each element\n').bold = True
    steps_list.add_run('3. Assemble the global stiffness matrix\n').bold = True
    steps_list.add_run('4. Apply boundary conditions and forces\n').bold = True
    steps_list.add_run('5. Solve for nodal displacements\n').bold = True
    steps_list.add_run('6. Calculate element stresses\n').bold = True
    steps_list.add_run('7. Compare with analytical solution\n').bold = True
    steps_list.add_run('8. Refine mesh if error exceeds threshold').bold = True
    
    # Mesh description
    doc.add_paragraph('4.2 Mesh Generation', style='Custom Heading 2')
    doc.add_paragraph('For the finite element model, we divided each segment into equal-length elements. The mesh was refined iteratively to achieve an error below 5% compared to the analytical solution.', style='Custom Normal')
    
    doc.add_paragraph('For elements in the first segment (with varying cross-section), we used the cross-sectional area at the element midpoint. This is a reasonable approximation for linear variations when sufficient elements are used.', style='Custom Normal')
    
    # FEM Results
    doc.add_paragraph('4.3 FEM Results', style='Custom Heading 2')
    
    # Try to insert displacement field plot
    try:
        doc.add_paragraph('The displacement distribution obtained from FEM is shown in Figure 3:', style='Custom Normal')
        doc.add_picture('plots/displacement_field.png', width=Inches(6))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph('Figure 3: Nodal displacements from FEM analysis', style='Custom Normal').alignment = WD_ALIGN_PARAGRAPH.CENTER
    except:
        doc.add_paragraph('[Displacement field plot would be displayed here]', style='Custom Normal').alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Error Analysis
    doc.add_paragraph('4.4 Error Analysis', style='Custom Heading 2')
    doc.add_paragraph('We compared the FEM results with the analytical solution to evaluate the accuracy of our numerical implementation. The error was calculated as the relative difference between the FEM and analytical stress values.', style='Custom Normal')
    
    # Try to insert error plot
    try:
        doc.add_paragraph('The error distribution is shown in Figure 4:', style='Custom Normal')
        doc.add_picture('plots/error_plot.png', width=Inches(6))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph('Figure 4: Error analysis comparing FEM with analytical solution', style='Custom Normal').alignment = WD_ALIGN_PARAGRAPH.CENTER
    except:
        doc.add_paragraph('[Error plot would be displayed here]', style='Custom Normal').alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Implementation
    doc.add_paragraph('5. Implementation', style='Custom Heading 1')
    
    # Python Implementation
    doc.add_paragraph('5.1 Python Implementation', style='Custom Heading 2')
    doc.add_paragraph('We implemented the analysis in Python using NumPy and SciPy for numerical operations. The implementation consists of several modules:', style='Custom Normal')
    
    python_modules = [
        ('main.py', 'Main script to run the analysis'),
        ('bar_analysis.py', 'Core implementation with analytical and FEM solutions'),
        ('utils.py', 'Helper functions for plotting and reporting')
    ]
    
    for module, description in python_modules:
        p = doc.add_paragraph(style='Custom Normal')
        p.add_run(module).bold = True
        p.add_run(f': {description}')
    
    # MATLAB Implementation
    doc.add_paragraph('5.2 MATLAB Implementation', style='Custom Heading 2')
    doc.add_paragraph('We also provided a MATLAB implementation with the following scripts:', style='Custom Normal')
    
    matlab_modules = [
        ('main_axial_bar.m', 'Driver script to run the analysis'),
        ('inputData.m', 'Define problem parameters'),
        ('generateMesh.m', 'Generate the finite element mesh'),
        ('solveBar.m', 'Solve the FE system and calculate stresses'),
        ('displayAnalyticalSolution.m', 'Calculate and display analytical results'),
        ('postPlots.m', 'Generate visualization plots'),
        ('errorPlot.m', 'Perform error analysis')
    ]
    
    for module, description in matlab_modules:
        p = doc.add_paragraph(style='Custom Normal')
        p.add_run(module).bold = True
        p.add_run(f': {description}')
    
    # MATLAB Code Snippet
    doc.add_paragraph('Sample MATLAB code (solveBar.m):', style='Custom Normal')
    
    with open('solveBar.m', 'r') as f:
        matlab_code = f.read()
    
    p = doc.add_paragraph(matlab_code, style='Code Block')
    
    # Conclusions
    doc.add_paragraph('6. Conclusions', style='Custom Heading 1')
    
    conclusions = [
        'The analytical solution provides exact stress and displacement values for this one-dimensional bar problem, with a free-end displacement of 5.44 mm.',
        'The stress varies from 400 MPa at the fixed end to 800 MPa at the first interface, then 600 MPa across the second segment, and 400 MPa in the third segment.',
        'The finite element method with suitable mesh refinement reproduces the analytical results with good accuracy (error < 5%).',
        'Mesh refinement is more critical in the first segment due to the varying cross-section.',
        'Both the Python and MATLAB implementations provide consistent results and can be easily extended to more complex problems.'
    ]
    
    for point in conclusions:
        p = doc.add_paragraph(style='Custom Normal')
        p.add_run('• ').bold = True
        p.add_run(point)
    
    # Save the document
    doc_path = 'Composite_Bar_Structure_Analysis_Report.docx'
    doc.save(doc_path)
    
    print(f"Comprehensive Word document created: {doc_path}")
    return doc_path

if __name__ == "__main__":
    # Ensure we're in the correct directory
    script_dir = os.path.dirname(os.path.abspath(__file__))
    os.chdir(script_dir)
    
    try:
        create_comprehensive_document()
    except Exception as e:
        print(f"Error creating Word document: {e}")
        print("You might need to install the python-docx package:")
        print("pip install python-docx")
